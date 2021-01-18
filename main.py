"""
python3
Este programa permitira escribir un documento de texto solo con la consola de python,
más especificamente una carta
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import smtplib

doc = Document()

"""Date and letter for letter"""
letter_date_input = input('Enter the date and place for your letter: ')
letter_date = doc.add_paragraph(letter_date_input)
letter_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

"""Regard for letter"""
regard_letter_input = input('Enter a greeting for the recipient: ')
regard_letter = doc.add_paragraph(regard_letter_input)
regard_letter.alignment = WD_ALIGN_PARAGRAPH.LEFT

"""Body for letter"""
body_input = input('Enter what you want to say: ')
body = doc.add_paragraph(body_input)
body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

"""Good bye for letter"""
gd_input = input('Enter the farewell: ')
gb = doc.add_paragraph(gd_input)
body.alignment = WD_ALIGN_PARAGRAPH.LEFT

"""Name"""
name_input = input('Enter your name: ')
name = doc.add_paragraph(name_input)
name.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.save('letter.docx')
